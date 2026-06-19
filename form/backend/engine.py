"""FormFiller engine — form detection + filling for PDF and DOCX.
Deterministic, no network. Used by the web backend."""
import re
from io import BytesIO

SYN = {
 'name':['full name','name','applicant name','applicant','deponent','party name','claimant','petitioner','plaintiff','respondent','defendant'],
 'date':['date','dated','date signed'],'dob':['date of birth','dob'],
 'address':['residential address','address','postal address','home address','mailing address','street address','address for service'],
 'phone':['phone number','phone','telephone','mobile','tel','contact number'],
 'email':['email','e-mail','email address'],
 'caseno':['case number','case no','court file number','suit no','cause no','docket','file number','reference'],
 'court':['name of court','court station','court','registry','court name'],
 'occupation':['occupation','profession','job title'],
 'authority':['statutory authority','authority','statute','provision','grounds','cause of action','order requested'],
 'signature':['signature','signed','sign'],
}
def norm(s): return re.sub(r'[^a-z0-9]','',(s or '').lower())
def synkey(label):
    n=norm(label)
    for k,vs in SYN.items():
        for v in vs:
            if n==norm(v) or n.startswith(norm(v)) or (len(norm(v))>4 and norm(v) in n): return k
    return None
def looks_legal(s): return bool(re.search(r'(statut|section\b|rule\b|regulation|\bcode\b|authority|cause of action|ground|cite|citation|\bcase\b|provision|order(s)? requested)',(s or ''),re.I))

# ---------- DETECT ----------
def detect(path):
    p=path.lower()
    if p.endswith('.docx'): return _detect_docx(path)
    if p.endswith('.pdf'): return _detect_pdf(path)
    raise ValueError('Unsupported file type (PDF or DOCX only).')

def _detect_pdf(path):
    import fitz
    doc=fitz.open(path); fields=[]; has=False
    for page in doc:
        words=page.get_text("words")
        for w in page.widgets() or []:
            has=True
            label=_nearest(words,w.rect) or (w.field_name or '')
            fields.append({'key':w.field_name or label,'label':label or (w.field_name or ''),
                           'kind':'check' if w.field_type==fitz.PDF_WIDGET_TYPE_CHECKBOX else 'text',
                           'legal':looks_legal(label)})
    if has: return {'kind':'pdf-fields','fields':fields}
    return {'kind':'pdf-flat','fields':[]}

def _nearest(words,rect):
    fx0,fy0,fx1,fy1=rect; fcy=(fy0+fy1)/2; best=None; bestd=1e9
    for w in words:
        x0,y0,x1,y1,word=w[0],w[1],w[2],w[3],w[4]; cy=(y0+y1)/2
        if abs(cy-fcy)<=max(fy1-fy0,y1-y0)*1.3 and x1<=fx0+2:
            d=fx0-x1
            if -40<d<bestd: bestd=d; best=word
    return (best or '').replace('_','').strip(' :')

def _detect_docx(path):
    from docx import Document
    doc=Document(path); fields=[]; seen=set()
    for p in doc.paragraphs:
        for tok in re.findall(r'(\{\{[^{}]+\}\}|\[[^\[\]\n]{1,60}?\]|\{[^{}]+\})', p.text):
            lab=re.sub(r'[{}\[\]]','',tok).strip()
            if tok not in seen: seen.add(tok); fields.append({'key':tok,'label':lab,'kind':'text','mode':'token','legal':looks_legal(lab)})
        m=re.match(r'^([A-Za-z][A-Za-z0-9 ._\-/()]{1,50}?):\s*(_{2,})?\s*$', p.text.strip())
        if m:
            lab=m.group(1).strip(); k='LBL:'+lab
            if k not in seen: seen.add(k); fields.append({'key':lab,'label':lab,'kind':'text','mode':'label','legal':looks_legal(lab)})
    for tb in doc.tables:
        for row in tb.rows:
            cells=row.cells
            for i,c in enumerate(cells):
                lab=c.text.strip()
                if lab and i+1<len(cells) and not cells[i+1].text.strip():
                    k='CELL:%d'%id(cells[i+1]._tc)
                    if k not in seen and len(lab)<60:
                        seen.add(k); fields.append({'key':k,'label':lab,'kind':'text','mode':'cell','legal':looks_legal(lab),
                                                    '_t':_tbidx(doc,tb),'_r':_rowidx(tb,row),'_c':i+1})
    return {'kind':'docx','fields':fields}

def _tbidx(doc,tb):
    for i,t in enumerate(doc.tables):
        if t is tb: return i
    return 0
def _rowidx(tb,row):
    for i,r in enumerate(tb.rows):
        if r is row: return i
    return 0

# ---------- FILL ----------
def fill(path, values, out_path):
    """values: dict key-> string (keys are field 'key' from detect, or labels)."""
    p=path.lower()
    if p.endswith('.docx'): return _fill_docx(path, values, out_path)
    return _fill_pdf(path, values, out_path)

def _resolve(label, values):
    n=norm(label)
    # exact key
    if label in values and values[label]!='' : return values[label]
    # normalized
    for k,v in values.items():
        if norm(k)==n and v!='' : return v
    sk=synkey(label)
    if sk:
        for k,v in values.items():
            if v!='' and synkey(k)==sk: return v
    for k,v in values.items():
        if v!='' and len(norm(k))>3 and (n in norm(k) or norm(k) in n): return v
    return None

def _fill_pdf(path, values, out_path):
    import fitz
    doc=fitz.open(path); rep=[]
    for page in doc:
        words=page.get_text("words")
        for w in page.widgets() or []:
            label=_nearest(words,w.rect) or (w.field_name or '')
            v = values.get(w.field_name) if w.field_name in values else _resolve(label, values)
            if v in (None,''): continue
            try:
                if w.field_type==fitz.PDF_WIDGET_TYPE_CHECKBOX:
                    w.field_value=bool(re.match(r'^(y|yes|true|x|on|1|check)',str(v),re.I))
                else: w.field_value=str(v)
                w.update(); rep.append((label or w.field_name, v))
            except Exception: pass
    doc.save(out_path, deflate=True); return rep

def _set(cell, text):
    cell.text=''
    parts=str(text).split('\n')
    for i,ln in enumerate(parts):
        (cell.paragraphs[0] if i==0 else cell.add_paragraph()).add_run(ln)

def _fill_docx(path, values, out_path):
    from docx import Document
    doc=Document(path); rep=[]
    for p in doc.paragraphs:
        t=p.text
        if not t.strip(): continue
        new=t
        for tok in re.findall(r'(\{\{[^{}]+\}\}|\[[^\[\]]+\]|\{[^{}]+\})', t):
            v=values.get(tok) or _resolve(re.sub(r'[{}\[\]]','',tok).strip(), values)
            if v: new=new.replace(tok,str(v)); rep.append((tok,v))
        m=re.match(r'^(?P<lab>[^:]{2,60}):\s*(?P<rest>.*)$', new)
        if m:
            v=_resolve(m.group('lab').strip(), values)
            if v and (re.search(r'_{2,}',m.group('rest')) or m.group('rest').strip()==''):
                new=f"{m.group('lab').strip()}: {v}"; rep.append((m.group('lab').strip(),v))
        if new!=t and (p.runs):
            p.runs[0].text=new
            for r in p.runs[1:]: r.text=''
        elif new!=t:
            p.add_run(new)
    # table cells (label -> adjacent)
    for tb in doc.tables:
        for row in tb.rows:
            cells=row.cells
            for i,c in enumerate(cells):
                lab=c.text.strip()
                if lab and i+1<len(cells) and not cells[i+1].text.strip():
                    v=_resolve(lab, values)
                    if v: _set(cells[i+1], v); rep.append((lab,v))
    doc.save(out_path); return rep
